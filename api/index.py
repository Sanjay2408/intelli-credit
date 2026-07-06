"""
IntelliCredit API — stateless FastAPI backend designed for serverless hosting.

Design notes:
- No vector DB / embedding model. Documents are parsed & chunked server-side,
  the chunks are returned to the client, which stores them and sends the
  relevant set back with every analysis request. Retrieval is BM25 (pure
  Python) computed per-request — deterministic, fast, and survives cold starts.
- Multi-format ingestion: PDF, CSV, TSV, TXT, MD, JSON, XLSX/XLS, DOCX, and
  images (PNG/JPG/WEBP via Groq vision OCR).
- All LLM calls go through Groq. Research uses groq/compound-mini which has
  built-in web search (no Serper key required).
"""
import io
import json
import math
import os
import re
import base64
from collections import Counter
from typing import Any, Dict, List, Optional

from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from pydantic import BaseModel

try:
    from dotenv import load_dotenv
    load_dotenv(os.path.join(os.path.dirname(__file__), "..", ".env"))
except Exception:
    pass

from groq import Groq

GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")
GROQ_MODEL = os.getenv("GROQ_MODEL", "llama-3.3-70b-versatile")
GROQ_VISION_MODEL = os.getenv("GROQ_VISION_MODEL", "meta-llama/llama-4-scout-17b-16e-instruct")
GROQ_RESEARCH_MODEL = os.getenv("GROQ_RESEARCH_MODEL", "groq/compound-mini")

MAX_CONTEXT_CHARS = 14000  # ~3.5k tokens of retrieved context per LLM call

app = FastAPI(title="IntelliCredit API", version="2.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


def groq_client() -> Groq:
    if not GROQ_API_KEY:
        raise HTTPException(500, "GROQ_API_KEY is not configured on the server.")
    return Groq(api_key=GROQ_API_KEY)


def llm(prompt: str, *, system: str = "", max_tokens: int = 1200,
        temperature: float = 0.2, json_mode: bool = False, model: str = "") -> str:
    client = groq_client()
    messages = []
    if system:
        messages.append({"role": "system", "content": system})
    messages.append({"role": "user", "content": prompt})
    kwargs: Dict[str, Any] = dict(
        model=model or GROQ_MODEL,
        messages=messages,
        temperature=temperature,
        max_tokens=max_tokens,
    )
    if json_mode:
        kwargs["response_format"] = {"type": "json_object"}
    last_err: Optional[Exception] = None
    for _ in range(3):
        try:
            resp = client.chat.completions.create(**kwargs)
            return resp.choices[0].message.content.strip()
        except Exception as exc:  # rate limits / transient errors
            last_err = exc
    raise HTTPException(502, f"Groq API error: {last_err}")


def llm_json(prompt: str, **kwargs) -> Dict:
    raw = llm(prompt, json_mode=True, **kwargs)
    raw = re.sub(r"^```[a-zA-Z]*\n?|\n?```$", "", raw.strip())
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        m = re.search(r"\{.*\}", raw, re.DOTALL)
        if m:
            return json.loads(m.group(0))
        raise HTTPException(502, "LLM returned malformed JSON.")


# ---------------------------------------------------------------------------
# Text extraction — every format
# ---------------------------------------------------------------------------

def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    pages = [(page.extract_text() or "") for page in reader.pages]
    text = "\n".join(p for p in pages if p.strip())
    if not text.strip():
        raise HTTPException(422, "No extractable text in this PDF (it may be scanned — upload as an image instead).")
    return text


def _extract_xlsx(data: bytes) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    out = []
    for ws in wb.worksheets:
        out.append(f"## Sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                out.append(" | ".join(cells))
    wb.close()
    return "\n".join(out)


def _extract_docx(data: bytes) -> str:
    import docx
    d = docx.Document(io.BytesIO(data))
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_csv(data: bytes, delimiter: str = ",") -> str:
    import csv
    text = data.decode("utf-8", errors="replace")
    rows = list(csv.reader(io.StringIO(text), delimiter=delimiter))
    if not rows:
        return ""
    header = rows[0]
    out = [" | ".join(header)]
    for row in rows[1:]:
        # Pair values with headers so chunks stay self-describing
        pairs = [f"{h}: {v}" for h, v in zip(header, row) if v.strip()]
        if pairs:
            out.append("; ".join(pairs))
    return "\n".join(out)


def _extract_image(data: bytes, mime: str) -> str:
    b64 = base64.b64encode(data).decode()
    client = groq_client()
    resp = client.chat.completions.create(
        model=GROQ_VISION_MODEL,
        messages=[{
            "role": "user",
            "content": [
                {"type": "text", "text": (
                    "Transcribe ALL text, numbers, and tables in this document image. "
                    "Preserve every financial figure exactly. Format tables as "
                    "'label: value' lines. Output only the transcription."
                )},
                {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}},
            ],
        }],
        temperature=0.0,
        max_tokens=4000,
    )
    return resp.choices[0].message.content.strip()


IMAGE_EXTS = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg", "webp": "image/webp"}


def extract_text(filename: str, data: bytes) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext == "pdf":
        return _extract_pdf(data)
    if ext in ("xlsx", "xlsm", "xls"):
        return _extract_xlsx(data)
    if ext == "docx":
        return _extract_docx(data)
    if ext == "csv":
        return _extract_csv(data)
    if ext == "tsv":
        return _extract_csv(data, delimiter="\t")
    if ext in IMAGE_EXTS:
        return _extract_image(data, IMAGE_EXTS[ext])
    if ext == "json":
        try:
            obj = json.loads(data.decode("utf-8", errors="replace"))
            return json.dumps(obj, indent=1, ensure_ascii=False)
        except json.JSONDecodeError:
            return data.decode("utf-8", errors="replace")
    if ext in ("txt", "md", "log", "text", ""):
        return data.decode("utf-8", errors="replace")
    # Last resort: try to decode as text
    text = data.decode("utf-8", errors="replace")
    if text.count("�") > len(text) * 0.1:
        raise HTTPException(415, f"Unsupported file type: .{ext}")
    return text


# ---------------------------------------------------------------------------
# Chunking + BM25 retrieval
# ---------------------------------------------------------------------------

def chunk_text(text: str, chunk_chars: int = 1600, overlap: int = 200) -> List[Dict]:
    """Split on paragraph boundaries into ~chunk_chars segments with overlap."""
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if not text:
        return []
    paras = re.split(r"\n\n+", text)
    chunks, buf = [], ""
    for para in paras:
        # Hard-split any paragraph longer than the chunk budget
        while len(para) > chunk_chars:
            head, para = para[:chunk_chars], para[chunk_chars - overlap:]
            chunks.append((buf + "\n\n" + head).strip() if buf else head)
            buf = ""
        if len(buf) + len(para) + 2 > chunk_chars and buf:
            chunks.append(buf.strip())
            buf = buf[-overlap:] + "\n\n" + para
        else:
            buf = (buf + "\n\n" + para) if buf else para
    if buf.strip():
        chunks.append(buf.strip())
    result = []
    for i, c in enumerate(chunks, 1):
        preview = " ".join(c.split()[:8])[:60]
        result.append({"section": f"Segment {i}: {preview}", "content": c})
    return result


_TOKEN_RE = re.compile(r"[a-z0-9]+")


def _tokenize(s: str) -> List[str]:
    return _TOKEN_RE.findall(s.lower())


def bm25_rank(query: str, chunks: List[Dict], top_k: int = 8) -> List[Dict]:
    """Score chunks against the query with BM25 (k1=1.5, b=0.75)."""
    if not chunks:
        return []
    q_terms = _tokenize(query)
    docs = [_tokenize(c.get("content", "")) for c in chunks]
    N = len(docs)
    avgdl = sum(len(d) for d in docs) / max(N, 1) or 1
    df: Counter = Counter()
    for d in docs:
        for t in set(d):
            df[t] += 1
    k1, b = 1.5, 0.75
    scored = []
    for c, d in zip(chunks, docs):
        tf = Counter(d)
        score = 0.0
        for t in q_terms:
            if t not in tf:
                continue
            idf = math.log(1 + (N - df[t] + 0.5) / (df[t] + 0.5))
            score += idf * tf[t] * (k1 + 1) / (tf[t] + k1 * (1 - b + b * len(d) / avgdl))
        scored.append((score, c))
    scored.sort(key=lambda x: x[0], reverse=True)
    out = []
    for score, c in scored[:top_k]:
        cc = dict(c)
        cc["score"] = round(score, 3)
        out.append(cc)
    return out


def build_context(chunks: List[Dict], queries: List[str], per_query_k: int = 5,
                  max_chars: int = MAX_CONTEXT_CHARS) -> str:
    """Multi-query BM25 retrieval packed into a character budget."""
    seen, picked = set(), []
    for q in queries:
        for c in bm25_rank(q, chunks, top_k=per_query_k):
            key = c.get("content", "")[:120]
            if key not in seen:
                seen.add(key)
                picked.append(c)
    ctx, total = [], 0
    for c in picked:
        block = f"[{c.get('source', 'doc')} — {c.get('section', '')}]\n{c.get('content', '')}"
        if total + len(block) > max_chars:
            break
        ctx.append(block)
        total += len(block)
    if not ctx and picked:
        ctx = [picked[0].get("content", "")[:max_chars]]
    return "\n\n".join(ctx)


# ---------------------------------------------------------------------------
# Schemas
# ---------------------------------------------------------------------------

class Chunk(BaseModel):
    section: str = ""
    content: str
    source: str = ""
    doc_type: str = "general"


class QueryRequest(BaseModel):
    question: str
    chunks: List[Chunk]
    top_k: int = 8


class AssessRequest(BaseModel):
    company_name: str
    sector: str = ""
    requested_loan_cr: float = 0
    chunks: List[Chunk]
    primary_insights: str = ""
    annual_revenue_cr: Optional[float] = None
    net_profit_cr: Optional[float] = None
    total_debt_cr: Optional[float] = None
    gstin: Optional[str] = None


class ChunksOnly(BaseModel):
    chunks: List[Chunk]
    company_name: str = ""


class SwotRequest(BaseModel):
    company_name: str
    sector: str = ""
    chunks: List[Chunk]


class ResearchRequest(BaseModel):
    company_name: str
    sector: str = ""
    query: str = ""


class CamRequest(BaseModel):
    company_name: str
    sector: str = ""
    assessment: Dict[str, Any] = {}
    chunks: List[Chunk] = []


DOC_TYPES = ["annual", "alm", "shareholding", "borrowing", "portfolio", "gst", "bank", "general"]
DOC_TYPE_LABELS = {
    "annual": "Annual Report", "alm": "ALM / Liquidity Statement",
    "shareholding": "Shareholding Pattern", "borrowing": "Borrowing Profile",
    "portfolio": "Loan Portfolio", "gst": "GST Returns",
    "bank": "Bank Statement", "general": "General Document",
}


def _chunks_dicts(chunks: List[Chunk]) -> List[Dict]:
    return [c.model_dump() for c in chunks]


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

@app.get("/api/health")
def health():
    return {"status": "ok", "version": "2.0.0"}


@app.post("/api/documents/process")
async def process_document(file: UploadFile = File(...), doc_type: str = Form("auto")):
    data = await file.read()
    if len(data) > 4 * 1024 * 1024:
        raise HTTPException(413, "File too large — max 4 MB per file on this deployment.")
    filename = file.filename or "document"
    text = extract_text(filename, data)
    if not text.strip():
        raise HTTPException(422, f"No text could be extracted from {filename}.")

    chunks = chunk_text(text)

    # One small LLM call: classify doc type + detect company name together
    detected_type, company = "general", ""
    try:
        meta = llm_json(
            "You are a financial document analyst. From the excerpt below return JSON "
            '{"doc_type": <one of ' + json.dumps(DOC_TYPES) + '>, '
            '"company_name": "<primary company whose data is reported, or empty string>"}\n\n'
            f"Filename: {filename}\nExcerpt:\n{text[:4000]}",
            max_tokens=100, temperature=0.0,
        )
        if meta.get("doc_type") in DOC_TYPES:
            detected_type = meta["doc_type"]
        company = str(meta.get("company_name") or "").strip()
        if company.lower() == "unknown":
            company = ""
    except HTTPException:
        pass

    final_type = doc_type if doc_type in DOC_TYPES else detected_type
    for c in chunks:
        c["source"] = filename
        c["doc_type"] = final_type

    return {
        "status": "success",
        "filename": filename,
        "num_chunks": len(chunks),
        "char_count": len(text),
        "doc_type": final_type,
        "doc_type_label": DOC_TYPE_LABELS.get(final_type, "General Document"),
        "company_name": company,
        "chunks": chunks,
    }


@app.post("/api/query")
def query_documents(req: QueryRequest):
    chunks = _chunks_dicts(req.chunks)
    if not chunks:
        raise HTTPException(400, "No documents indexed. Upload documents first.")
    top = bm25_rank(req.question, chunks, top_k=req.top_k)
    context = build_context(chunks, [req.question], per_query_k=req.top_k)
    answer = llm(
        "You are a financial analyst AI for a credit appraisal system. Answer using ONLY "
        "the document context below. Be precise and factual; cite figures exactly. Use "
        "bullet points where helpful. If the answer is not in the context, say so clearly.\n\n"
        f"Context:\n{context}\n\nQuestion: {req.question}",
        max_tokens=900, temperature=0.1,
    )
    sources = [{"section": c.get("section", ""), "source": c.get("source", ""),
                "score": c.get("score", 0), "text": c.get("content", "")[:300]} for c in top[:5]]
    return {"answer": answer, "sources": sources}


@app.post("/api/assess/extract-form")
def extract_form(req: ChunksOnly):
    chunks = _chunks_dicts(req.chunks)
    if not chunks:
        raise HTTPException(400, "No documents indexed.")
    context = build_context(chunks, [
        "company name registered office", "total revenue turnover income",
        "net profit after tax", "total debt borrowings loans", "GSTIN GST number sector industry",
    ], per_query_k=4)
    data = llm_json(
        "Extract these fields from the financial document context. Return JSON:\n"
        '{"company_name": str, "sector": str, "annual_revenue_cr": number|null, '
        '"net_profit_cr": number|null, "total_debt_cr": number|null, "gstin": str|null}\n'
        "Amounts must be in INR Crores (convert lakhs/millions). Use null when unknown.\n\n"
        f"Context:\n{context}",
        max_tokens=250, temperature=0.0,
    )
    return data


@app.post("/api/assess")
def assess(req: AssessRequest):
    chunks = _chunks_dicts(req.chunks)
    if not chunks:
        raise HTTPException(400, "No documents indexed. Upload documents first.")
    context = build_context(chunks, [
        f"{req.company_name} revenue profit financial performance",
        "debt borrowings repayment interest liabilities",
        "GST compliance tax litigation regulatory",
        "management promoters shareholding collateral assets",
    ], per_query_k=4)

    knowns = []
    if req.annual_revenue_cr: knowns.append(f"Annual revenue: ₹{req.annual_revenue_cr} Cr")
    if req.net_profit_cr is not None and req.net_profit_cr != 0: knowns.append(f"Net profit: ₹{req.net_profit_cr} Cr")
    if req.total_debt_cr: knowns.append(f"Total debt: ₹{req.total_debt_cr} Cr")
    if req.gstin: knowns.append(f"GSTIN: {req.gstin}")

    prompt = (
        "You are a senior credit officer at an Indian bank preparing a corporate credit decision.\n\n"
        f"Company: {req.company_name}\nSector: {req.sector or 'Unknown'}\n"
        f"Requested loan: ₹{req.requested_loan_cr} Cr\n"
        + (("Analyst-provided figures: " + "; ".join(knowns) + "\n") if knowns else "")
        + (f"Primary due-diligence notes (site visits, interviews): {req.primary_insights}\n" if req.primary_insights else "")
        + f"\nDocument evidence:\n{context}\n\n"
        "Produce a rigorous, explainable credit assessment. Return JSON exactly:\n"
        "{\n"
        '  "risk_score": <0-100 int, higher = riskier>,\n'
        '  "risk_band": "Low"|"Medium"|"High",\n'
        '  "decision": "approved"|"conditional"|"rejected",\n'
        '  "recommended_loan_cr": <number>,\n'
        '  "interest_rate_pct": <number>,\n'
        '  "tenor_months": <int>,\n'
        '  "score_breakdown": {"financial_health": <0-100>, "repayment_history": <0-100>, '
        '"collateral_coverage": <0-100>, "management_quality": <0-100>, "market_position": <0-100>},\n'
        '  "financial_overview": {"annual_revenue_cr": <number|null>, "net_profit_cr": <number|null>, '
        '"total_debt_cr": <number|null>, "gst_turnover_cr": <number|null>},\n'
        '  "profitability_metrics": {"net_margin_pct": <number|null>, "roe_pct": <number|null>, '
        '"dscr": <number|null>, "debt_equity": <number|null>, "interest_coverage": <number|null>},\n'
        '  "yearly_trend": [{"year": "FY23", "revenue_cr": <num>, "profit_cr": <num>, "debt_cr": <num>}, ...],\n'
        '  "risk_alerts": [{"severity": "High"|"Medium"|"Low", "title": str, "detail": str}, ...],\n'
        '  "conditions": [str, ...],\n'
        '  "reasoning": "<3-4 paragraph narrative explaining the decision>"\n'
        "}\n"
        "Score breakdown weights: financial health 30%, repayment 25%, collateral 20%, "
        "management 15%, market 10%. Base every number on document evidence; use null "
        "when evidence is absent — do not invent figures. yearly_trend only if the "
        "documents contain multi-year data (else empty array)."
    )
    result = llm_json(prompt, max_tokens=2600, temperature=0.15)
    # Normalize risk band from score if the LLM omitted/contradicted it
    score = int(result.get("risk_score", 50))
    result["risk_score"] = max(0, min(100, score))
    result.setdefault("risk_band", "Low" if score < 35 else "Medium" if score < 65 else "High")
    return result


@app.post("/api/charts/financial")
def charts(req: ChunksOnly):
    chunks = _chunks_dicts(req.chunks)
    if not chunks:
        raise HTTPException(400, "No documents indexed.")
    context = build_context(chunks, [
        f"{req.company_name} revenue turnover yearly",
        "profit after tax net income yearly",
        "debt borrowings total liabilities", "margin ratio ROE DSCR interest coverage",
    ], per_query_k=4)
    return llm_json(
        "Extract financial chart data from the context. Amounts in INR Crores. Return JSON:\n"
        '{"financial_overview": {"annual_revenue_cr": num|null, "net_profit_cr": num|null, '
        '"total_debt_cr": num|null, "gst_turnover_cr": num|null},\n'
        '"yearly_trend": [{"year": str, "revenue_cr": num, "profit_cr": num, "debt_cr": num}],\n'
        '"profitability_metrics": {"gross_margin_pct": num|null, "net_margin_pct": num|null, '
        '"roe_pct": num|null, "dscr": num|null, "debt_equity": num|null, "interest_coverage": num|null}}\n'
        "Use null / empty array when the context lacks the data. Never invent numbers.\n\n"
        f"Context:\n{context}",
        max_tokens=700, temperature=0.0,
    )


@app.post("/api/swot")
def swot(req: SwotRequest):
    chunks = _chunks_dicts(req.chunks)
    if not chunks:
        raise HTTPException(400, "No documents indexed.")
    context = build_context(chunks, [
        f"{req.company_name} strengths market position growth",
        "weaknesses losses debt risks litigation",
        "opportunities expansion new markets", "threats competition regulation sector",
    ], per_query_k=4)
    return llm_json(
        f"Perform a sector-aware SWOT analysis of {req.company_name} "
        f"({req.sector or 'sector unknown'}) for a bank credit committee, grounded in the "
        "document evidence below. Return JSON:\n"
        '{"strengths": [str], "weaknesses": [str], "opportunities": [str], "threats": [str], '
        '"summary": "<2-3 sentence overall read>"}\n'
        "3-5 concise, specific items per quadrant citing actual figures where available.\n\n"
        f"Context:\n{context}",
        max_tokens=1100, temperature=0.2,
    )


@app.post("/api/gst/validate")
def gst_validate(req: ChunksOnly):
    chunks = _chunks_dicts(req.chunks)
    if not chunks:
        raise HTTPException(400, "No documents indexed.")
    context = build_context(chunks, [
        "GST GSTR turnover taxable value returns filed",
        "bank statement credits deposits receipts",
        "revenue turnover as per financial statements", "input tax credit ITC e-way bill",
    ], per_query_k=5)
    return llm_json(
        "You are a GST reconciliation specialist for Indian corporate credit appraisal. "
        "Cross-validate reported turnover across GST returns (GSTR-1/2A/3B), bank statement "
        "credits, and financial statements in the evidence below. Flag circular trading, "
        "revenue inflation, ITC mismatch, and delayed filings. Return JSON:\n"
        '{"status": "consistent"|"minor_mismatch"|"major_mismatch"|"insufficient_data",\n'
        '"gst_turnover_cr": num|null, "bank_credits_cr": num|null, "books_revenue_cr": num|null,\n'
        '"variance_pct": num|null,\n'
        '"findings": [{"severity": "High"|"Medium"|"Low", "title": str, "detail": str}],\n'
        '"recommendation": str}\n\n'
        f"Evidence:\n{context}",
        max_tokens=900, temperature=0.1,
    )


@app.post("/api/research")
def research(req: ResearchRequest):
    """Live web research via Groq compound (built-in web search)."""
    if req.query:
        prompt = req.query
    else:
        prompt = (
            f"Research the Indian company '{req.company_name}'"
            + (f" ({req.sector} sector)" if req.sector else "")
            + ". Cover: recent news, financial performance, litigation/regulatory actions "
            "(SEBI, RBI, MCA, courts), promoter background, and sector outlook. "
            "Summarize findings relevant to a bank assessing credit risk."
        )
    client = groq_client()
    try:
        resp = client.chat.completions.create(
            model=GROQ_RESEARCH_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
            max_tokens=1800,
        )
        raw = resp.choices[0].message.content.strip()
    except Exception as exc:
        raise HTTPException(502, f"Web research failed: {exc}")

    # Structure the raw research into severity-rated insights
    try:
        structured = llm_json(
            "Convert this research report into JSON: "
            '{"insights": [{"category": str, "severity": "High"|"Medium"|"Low", '
            '"title": str, "detail": str}], "summary": str}\n'
            "4-8 insights. Severity reflects credit-risk impact.\n\n" + raw[:9000],
            max_tokens=1200, temperature=0.1,
        )
    except HTTPException:
        structured = {"insights": [], "summary": raw[:1200]}
    structured["raw_report"] = raw
    return structured


@app.post("/api/cam/generate")
def cam_generate(req: CamRequest):
    chunks = _chunks_dicts(req.chunks)
    context = build_context(chunks, [
        f"{req.company_name} overview business operations",
        "financial ratios revenue profit debt", "collateral security management",
    ], per_query_k=3, max_chars=8000) if chunks else ""
    assessment_json = json.dumps(req.assessment)[:6000]
    return llm_json(
        "Draft a bank-grade Credit Appraisal Memo (CAM). Return JSON:\n"
        '{"sections": [{"title": str, "content": "<detailed markdown-free prose, 150-300 words>"}]}\n'
        "Sections, in order: Company Overview; Financial Analysis; 5Cs of Credit "
        "(Character, Capacity, Capital, Collateral, Conditions); Risk Assessment; "
        "Credit Decision & Terms; Compliance & Due Diligence.\n\n"
        f"Company: {req.company_name} | Sector: {req.sector}\n"
        f"Assessment result: {assessment_json}\n"
        + (f"Document evidence:\n{context}" if context else ""),
        max_tokens=2800, temperature=0.2,
    )


@app.post("/api/cam/docx")
def cam_docx(req: CamRequest):
    """Generate the CAM as a downloadable Word document."""
    sections = req.assessment.get("cam_sections") or []
    if not sections:
        cam = cam_generate(req)
        sections = cam.get("sections", [])
    import docx
    from docx.shared import Pt, RGBColor
    doc = docx.Document()
    title = doc.add_heading("Credit Appraisal Memo", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)
    meta = doc.add_paragraph()
    meta.add_run(f"Company: {req.company_name}    Sector: {req.sector or '—'}").bold = True
    a = req.assessment
    if a.get("risk_score") is not None:
        doc.add_paragraph(
            f"Risk Score: {a.get('risk_score')}/100 ({a.get('risk_band', '')})   "
            f"Decision: {str(a.get('decision', '')).title()}   "
            f"Recommended: ₹{a.get('recommended_loan_cr', '—')} Cr @ {a.get('interest_rate_pct', '—')}% "
            f"for {a.get('tenor_months', '—')} months"
        )
    for s in sections:
        doc.add_heading(s.get("title", "Section"), level=1)
        p = doc.add_paragraph(s.get("content", ""))
        for run in p.runs:
            run.font.size = Pt(10.5)
    buf = io.BytesIO()
    doc.save(buf)
    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="CAM_{re.sub(r"[^A-Za-z0-9]+", "_", req.company_name)}.docx"'},
    )
